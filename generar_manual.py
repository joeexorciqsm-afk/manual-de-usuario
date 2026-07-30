"""
Generador del Manual de Usuario de VORTEX  (documento bilingue)
VORTEX User Manual generator (bilingual document)

Sistema de Gestion para Taller de Bombas de Agua
Water Pump Workshop Management System

El contenido se define una sola vez, con cada cadena como par (espanol, ingles).
El documento se arma dos veces desde la misma estructura: primero la version en
espanol y despues la version en ingles, de modo que ambas son identicas en
organizacion, tablas y figuras.

Uso:
    pip install python-docx
    python3 generar_manual.py

Si existen ./capturas/figura1.png ... figura5.png se incrustan automaticamente.
Si no existen, se dibuja un marco reservado del mismo tamano.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ES, EN = 0, 1

# ---------------------------------------------------------------- paleta
NAVY = RGBColor(0x0F, 0x22, 0x33)
NAVY_HEX = "0F2233"
BLUE = RGBColor(0x15, 0x84, 0xD8)
BLUE_HEX = "1584D8"
GRAY = RGBColor(0x5A, 0x6A, 0x78)
INK = RGBColor(0x20, 0x2A, 0x33)
LIGHT_HEX = "EEF2F6"
ZEBRA_HEX = "F7F9FB"
BORDER_HEX = "C9D4DE"

BASE = os.path.dirname(os.path.abspath(__file__))
CAPTURAS = os.path.join(BASE, "capturas")

FIG_W = 6.3      # ancho de figura en pulgadas
FIG_H = 3.95     # alto equivalente para una captura de 1024x642


# ============================================================ utilidades
def set_cell_bg(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_borders(cell, hex_color=BORDER_HEX, sz=4):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:color"), hex_color)
        borders.append(el)
    tcPr.append(borders)


def add_field(paragraph, code, size=9):
    run = paragraph.add_run()
    f1 = OxmlElement("w:fldChar")
    f1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = code
    f2 = OxmlElement("w:fldChar")
    f2.set(qn("w:fldCharType"), "end")
    run._r.append(f1)
    run._r.append(instr)
    run._r.append(f2)
    run.font.size = Pt(size)
    run.font.color.rgb = GRAY
    return run


def restart_page_numbering(section, start=1):
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:start"), str(start))
    section._sectPr.append(pg)


def horizontal_rule(paragraph, hex_color=BLUE_HEX, sz=18):
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), hex_color)
    borders.append(bottom)
    pPr.append(borders)


_BM = {"id": 100}


def bookmark_start(paragraph, name):
    """Abre un marcador. Permite acotar cada indice a su idioma."""
    _BM["id"] += 1
    bid = str(_BM["id"])
    el = OxmlElement("w:bookmarkStart")
    el.set(qn("w:id"), bid)
    el.set(qn("w:name"), name)
    paragraph._p.insert(0, el)
    return bid


def bookmark_end(paragraph, bid):
    el = OxmlElement("w:bookmarkEnd")
    el.set(qn("w:id"), bid)
    paragraph._p.append(el)


def add_toc(doc, bookmark):
    """Indice automatico limitado al marcador indicado."""
    par = doc.add_paragraph()
    run = par.add_run()
    f1 = OxmlElement("w:fldChar")
    f1.set(qn("w:fldCharType"), "begin")
    f1.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f'TOC \\o "1-2" \\h \\z \\u \\b {bookmark}'
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = ("Seleccione esta tabla y presione F9 para generar el indice. / "
                "Select this table and press F9 to build the index.")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for e in (f1, instr, sep, txt, end):
        run._r.append(e)
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    return par


def para(doc, text="", size=10.5, bold=False, italic=False, color=None,
         align=None, space_after=6, space_before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        r.font.color.rgb = color if color else INK
    return p


def data_table(doc, rows, widths=None, font_size=9.5):
    tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
    tbl.style = "Table Grid"
    tbl.autofit = False
    for i, fila in enumerate(rows):
        for j, valor in enumerate(fila):
            cell = tbl.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(valor))
            r.font.size = Pt(font_size)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if i == 0:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                set_cell_bg(cell, NAVY_HEX)
            elif i % 2 == 0:
                set_cell_bg(cell, ZEBRA_HEX)
            if widths:
                cell.width = Inches(widths[j])
    return tbl


# --------------------------------------------------------------- figuras
def listar_capturas():
    if not os.path.isdir(CAPTURAS):
        return []
    val = (".png", ".jpg", ".jpeg")
    return sorted([f for f in os.listdir(CAPTURAS) if f.lower().endswith(val)],
                  key=lambda s: s.lower())


DISPONIBLES = listar_capturas()
MAPEO = {}


def buscar_captura(numero):
    for ext in ("png", "jpg", "jpeg", "PNG", "JPG", "JPEG"):
        cand = os.path.join(CAPTURAS, f"figura{numero}.{ext}")
        if os.path.exists(cand):
            return cand
    if len(DISPONIBLES) >= numero:
        return os.path.join(CAPTURAS, DISPONIBLES[numero - 1])
    return None


def figura(doc, numero, titulo, lang):
    ruta = buscar_captura(numero)
    MAPEO[numero] = os.path.basename(ruta) if ruta else "-- marco reservado --"
    etiqueta = "Figura" if lang == ES else "Figure"

    if ruta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.add_run().add_picture(ruta, width=Inches(FIG_W))
    else:
        tbl = doc.add_table(rows=1, cols=1)
        tbl.autofit = False
        cell = tbl.cell(0, 0)
        cell.width = Inches(FIG_W + 0.1)
        set_cell_bg(cell, LIGHT_HEX)
        set_cell_borders(cell, BORDER_HEX, sz=8)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row = tbl.rows[0]
        row.height = Inches(FIG_H)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p1.add_run(f"[ {etiqueta} {numero} ]")
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = GRAY
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(titulo)
        r2.italic = True
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = GRAY

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(f"{etiqueta} {numero}. ")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = NAVY
    r2 = cap.add_run(titulo)
    r2.font.size = Pt(9)
    r2.italic = True
    r2.font.color.rgb = GRAY


def nota(doc, texto, lang):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    cell = tbl.cell(0, 0)
    cell.width = Inches(FIG_W + 0.1)
    set_cell_bg(cell, LIGHT_HEX)
    set_cell_borders(cell, BLUE_HEX, sz=6)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(("Nota: " if lang == ES else "Note: "))
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = NAVY
    r2 = p.add_run(texto)
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ==========================================================================
#  CONTENIDO BILINGUE   ( "espanol", "english" )
#  Los nombres de botones y campos se conservan en espanol, tal como
#  aparecen en pantalla, con su traduccion entre parentesis.
# ==========================================================================
T = lambda es, en: (es, en)

BLOQUES = [

    # ---------------------------------------------------- 1. INTRODUCCION
    ("h1", T("1. Introduccion", "1. Introduction")),
    ("p", T(
        "VORTEX es el sistema de gestion de un taller de bombas de agua. Reune en "
        "una sola aplicacion la recepcion de solicitudes de servicio de los "
        "clientes, el seguimiento de esas solicitudes, el control del inventario "
        "de piezas y la consulta de reportes de ventas.",
        "VORTEX is the management system for a water pump workshop. It brings "
        "together in a single application the intake of customer service requests, "
        "the tracking of those requests, parts inventory control and access to "
        "sales reports.")),
    ("p", T("El sistema se organiza en dos entornos de trabajo independientes:",
            "The system is organized into two independent working environments:")),
    ("bullets", [
        {"b": T("Modulo Cliente.", "Client Module."),
         "t": T("Pantalla publica donde una persona registra su solicitud de "
                "servicio sin necesidad de tener cuenta ni contrasena.",
                "A public screen where a person submits a service request with no "
                "account or password required.")},
        {"b": T("Modulo Administrador.", "Administrator Module."),
         "t": T("Area privada, protegida por inicio de sesion, donde el personal "
                "del taller gestiona solicitudes, inventario y reportes.",
                "A private area, protected by login, where workshop staff manage "
                "requests, inventory and reports.")},
    ]),
    ("p", T(
        "Este manual describe el funcionamiento de cada pantalla, el significado "
        "de sus datos y las acciones disponibles en cada una.",
        "This manual describes how each screen works, what its data means and "
        "which actions are available on each one.")),

    ("h2", T("1.1 Perfiles de usuario", "1.1 User profiles")),
    ("table", [
        [T("Perfil", "Profile"), T("Acceso", "Access"),
         T("Que puede hacer", "What they can do")],
        [T("Cliente", "Client"), T("Libre, sin credenciales", "Open, no credentials"),
         T("Registrar una solicitud de servicio y obtener su numero de pedido.",
           "Submit a service request and obtain an order number.")],
        [T("Administrador", "Administrator"),
         T("Requiere iniciar sesion", "Requires login"),
         T("Consultar el resumen del taller, administrar clientes y solicitudes, "
           "mantener el inventario de piezas y revisar reportes de ventas.",
           "Review the workshop summary, manage clients and requests, maintain the "
           "parts inventory and review sales reports.")],
    ], [1.2, 1.6, 3.6]),

    # ------------------------------------------------------ 2. CLIENTE
    ("pagebreak",),
    ("h1", T("2. Modulo Cliente: Solicitar servicio",
             "2. Client Module: Request a service")),
    ("p", T(
        "Es la pantalla con la que inicia el sistema. Su unico proposito es "
        "capturar la solicitud de servicio de un cliente. La pantalla se divide en "
        "un menu lateral oscuro a la izquierda y el formulario de captura a la "
        "derecha.",
        "This is the screen the system starts on. Its only purpose is to capture a "
        "client's service request. The screen is split into a dark side menu on the "
        "left and the entry form on the right.")),
    ("figura", 1, T("Pantalla Solicitar servicio del modulo Cliente",
                    "The Solicitar servicio (Request service) screen of the Client Module")),

    ("h2", T("2.1 Menu lateral", "2.1 Side menu")),
    ("p", T("La barra lateral identifica el modulo activo y concentra la navegacion:",
            "The side bar identifies the active module and holds the navigation:")),
    ("bullets", [
        {"b": T("VORTEX / Cliente.", "VORTEX / Cliente."),
         "t": T("Encabezado que confirma que se esta trabajando en el entorno de "
                "cliente.",
                "A header confirming that you are working in the client "
                "environment.")},
        {"b": T("Solicitar servicio.", "Solicitar servicio (Request service)."),
         "t": T("Unica opcion del menu y por eso aparece siempre resaltada en azul: "
                "indica la seccion en la que se encuentra.",
                "The only menu option, which is why it always appears highlighted "
                "in blue: it marks the section you are in.")},
        {"b": T("Iniciar sesion como admin.",
                "Iniciar sesion como admin (Sign in as admin)."),
         "t": T("Boton al pie del menu, precedido por la pregunta \u00bfEres admin? "
                "Da paso a la pantalla de acceso del personal del taller.",
                "A button at the foot of the menu, preceded by the question "
                "\u00bfEres admin? (Are you an admin?). It leads to the login "
                "screen for workshop staff.")},
    ]),

    ("h2", T("2.2 Campos del formulario", "2.2 Form fields")),
    ("p", T(
        "El texto guia bajo el titulo indica al cliente que debe llenar sus datos y "
        "describir el problema, y le anticipa que al enviar recibira un numero de "
        "pedido. Los campos solicitados son:",
        "The helper text under the title tells the client to fill in their details "
        "and describe the problem, and lets them know that on submitting they will "
        "receive an order number. The requested fields are:")),
    ("table", [
        [T("Campo", "Field"), T("Tipo", "Type"),
         T("Que se registra", "What is recorded")],
        [T("Nombre", "Nombre (Name)"), T("Linea de texto", "Single-line text"),
         T("Nombre de la persona o del negocio que solicita el servicio. Es el "
           "nombre con el que la solicitud aparecera en el modulo de administracion.",
           "Name of the person or business requesting the service. This is the name "
           "the request will appear under in the administration module.")],
        [T("Numero (telefono)", "Numero (telefono) (Phone number)"),
         T("Linea de texto", "Single-line text"),
         T("Telefono de contacto para confirmar la visita o avisar del avance.",
           "Contact phone number used to confirm the visit or report progress.")],
        [T("Direccion", "Direccion (Address)"),
         T("Linea de texto", "Single-line text"),
         T("Domicilio donde se prestara el servicio; orienta al tecnico que sera "
           "asignado.",
           "Address where the service will be performed; it guides the technician "
           "who will be assigned.")],
        [T("Motivo por el que solicita el servicio",
           "Motivo por el que solicita el servicio (Reason for the request)"),
         T("Area de texto amplia", "Large text area"),
         T("Descripcion del problema de la bomba. Al ser un campo grande admite "
           "varias lineas, de modo que el cliente puede detallar sintomas, "
           "antecedentes o el modelo del equipo.",
           "Description of the pump problem. Being a large field it accepts several "
           "lines, so the client can detail symptoms, history or the equipment "
           "model.")],
    ], [1.7, 1.1, 3.6]),

    ("h2", T("2.3 Envio de la solicitud", "2.3 Submitting the request")),
    ("p", T("Para registrar una solicitud:", "To submit a request:")),
    ("steps", [
        T("Escribir el nombre completo en el campo Nombre.",
          "Type the full name in the Nombre (Name) field."),
        T("Capturar un telefono de contacto valido en Numero (telefono).",
          "Enter a valid contact phone number in Numero (telefono)."),
        T("Indicar la direccion donde se encuentra el equipo en Direccion.",
          "Give the address where the equipment is located in Direccion."),
        T("Describir la falla en el area Motivo por el que solicita el servicio.",
          "Describe the fault in the Motivo por el que solicita el servicio area."),
        T("Pulsar el boton Enviar solicitud.",
          "Click the Enviar solicitud (Submit request) button."),
    ]),
    ("p", T(
        "Al confirmar el envio el sistema genera un numero de pedido, que es el "
        "identificador con el que el cliente puede dar seguimiento a su servicio. "
        "Ese mismo consecutivo es el que el administrador vera despues en la "
        "columna ID de las pantallas de gestion.",
        "Once the submission is confirmed the system generates an order number, "
        "which is the identifier the client uses to follow up on their service. "
        "That same sequential number is what the administrator will later see in "
        "the ID column of the management screens.")),
    ("nota", T(
        "La solicitud entra al sistema con el estado PENDIENTE. A partir de ese "
        "momento queda visible para el administrador tanto en el panel de Inicio "
        "como en la pantalla Clientes y solicitudes.",
        "The request enters the system with the status PENDIENTE (pending). From "
        "that moment it is visible to the administrator both on the Inicio (Home) "
        "dashboard and on the Clientes y solicitudes screen.")),

    # -------------------------------------------------- 3. ADMINISTRADOR
    ("pagebreak",),
    ("h1", T("3. Modulo Administrador", "3. Administrator Module")),
    ("p", T(
        "Tras iniciar sesion se abre el area de administracion. Todas sus pantallas "
        "comparten el mismo menu lateral, lo que permite cambiar de seccion sin "
        "perder el contexto de trabajo.",
        "After logging in, the administration area opens. All of its screens share "
        "the same side menu, which makes it possible to switch sections without "
        "losing the working context.")),

    ("h2", T("3.1 Menu lateral del administrador", "3.1 Administrator side menu")),
    ("p", T(
        "El menu muestra el encabezado VORTEX con la leyenda Administrador y cuatro "
        "opciones de navegacion. La opcion activa se resalta en azul:",
        "The menu shows the VORTEX header with the label Administrador and four "
        "navigation options. The active option is highlighted in blue:")),
    ("table", [
        [T("Opcion", "Option"),
         T("Pantalla a la que conduce", "Screen it leads to")],
        [T("Inicio", "Inicio (Home)"),
         T("Panel de resumen con los indicadores generales del taller y los "
           "proximos servicios pendientes.",
           "Summary dashboard with the workshop's general indicators and the "
           "upcoming pending services.")],
        [T("Clientes", "Clientes (Clients)"),
         T("Listado completo de clientes y solicitudes, con las acciones de gestion "
           "sobre cada una.",
           "Full list of clients and requests, with the management actions "
           "available for each one.")],
        [T("Inventario / Piezas", "Inventario / Piezas (Inventory / Parts)"),
         T("Catalogo de piezas con precios y existencias.",
           "Parts catalog with prices and stock levels.")],
        [T("Reportes", "Reportes (Reports)"),
         T("Estadisticas de ventas, piezas mas vendidas y alertas de stock bajo.",
           "Sales statistics, best-selling parts and low stock alerts.")],
    ], [1.7, 4.7]),
    ("p", T(
        "En la parte inferior el menu identifica al usuario conectado como "
        "Administrador e incluye el boton Cerrar sesion, que finaliza la sesion y "
        "devuelve el sistema a la pantalla del cliente.",
        "At the bottom, the menu identifies the signed-in user as Administrador and "
        "includes the Cerrar sesion (Sign out) button, which ends the session and "
        "returns the system to the client screen.")),

    # ------------------------------------------------------- 3.2 INICIO
    ("pagebreak",),
    ("h2", T("3.2 Inicio: panel de resumen", "3.2 Home: summary dashboard")),
    ("p", T(
        "Es la pantalla de bienvenida del administrador. Presenta el estado general "
        "del taller en un solo vistazo mediante cuatro indicadores y una tabla de "
        "servicios proximos.",
        "This is the administrator's welcome screen. It presents the overall state "
        "of the workshop at a glance through four indicators and a table of "
        "upcoming services.")),
    ("figura", 2, T("Panel de Inicio del modulo Administrador",
                    "The Inicio (Home) dashboard of the Administrator Module")),

    ("h3", T("3.2.1 Indicadores generales", "3.2.1 General indicators")),
    ("p", T(
        "Las cuatro tarjetas superiores resumen la operacion. Cada una se distingue "
        "por el color de su franja lateral:",
        "The four cards at the top summarize operations. Each is distinguished by "
        "the color of its side stripe:")),
    ("table", [
        [T("Indicador", "Indicator"), T("Significado", "Meaning"),
         T("Ejemplo", "Example")],
        [T("Clientes registrados", "Clientes registrados (Registered clients)"),
         T("Numero total de clientes dados de alta en el sistema.",
           "Total number of clients registered in the system."), T("30", "30")],
        [T("Tipos de pieza", "Tipos de pieza (Part types)"),
         T("Cantidad de piezas distintas del catalogo y, entre parentesis, la suma "
           "de todas las existencias.",
           "Number of distinct parts in the catalog and, in parentheses, the sum of "
           "all stock on hand."),
         T("149 (3311 en stock)", "149 (3311 in stock)")],
        [T("Ventas realizadas", "Ventas realizadas (Sales completed)"),
         T("Total de ventas registradas historicamente.",
           "Total number of sales recorded historically."), T("12", "12")],
        [T("Ingresos totales", "Ingresos totales (Total revenue)"),
         T("Suma del importe de todas las ventas registradas.",
           "Sum of the amount of every recorded sale."),
         T("$16,735.00", "$16,735.00")],
    ], [1.5, 3.6, 1.3]),
    ("p", T(
        "La lectura conjunta de estas tarjetas permite evaluar el tamano de la "
        "cartera de clientes, la profundidad del inventario y el desempeno "
        "comercial acumulado.",
        "Read together, these cards make it possible to assess the size of the "
        "client base, the depth of the inventory and cumulative commercial "
        "performance.")),

    ("h3", T("3.2.2 Proximos servicios (pendientes)",
             "3.2.2 Upcoming services (pending)")),
    ("p", T(
        "Debajo de los indicadores se muestra la tabla Proximos servicios, cuyo "
        "titulo incluye el numero de solicitudes pendientes de atender. Solo lista "
        "las solicitudes que aun no se han resuelto, ordenadas por fecha y hora, de "
        "manera que la primera fila corresponde al servicio mas inmediato. Sus "
        "columnas son:",
        "Below the indicators is the Proximos servicios (Upcoming services) table, "
        "whose title includes the number of requests still to be handled. It lists "
        "only unresolved requests, sorted by date and time, so the first row is the "
        "most immediate service. Its columns are:")),
    ("table", [
        [T("Columna", "Column"), T("Contenido", "Content")],
        [T("ID", "ID"),
         T("Numero de pedido de la solicitud, el mismo que recibio el cliente al "
           "enviarla.",
           "Order number of the request, the same one the client received on "
           "submitting it.")],
        [T("Cliente", "Cliente (Client)"),
         T("Nombre de la persona o negocio solicitante.",
           "Name of the requesting person or business.")],
        [T("Telefono", "Telefono (Phone)"),
         T("Numero de contacto capturado en la solicitud.",
           "Contact number entered in the request.")],
        [T("Direccion", "Direccion (Address)"),
         T("Domicilio donde debe prestarse el servicio.",
           "Address where the service is to be performed.")],
        [T("Fecha", "Fecha (Date)"),
         T("Dia programado o de registro del servicio.",
           "Scheduled or recorded day of the service.")],
        [T("Hora", "Hora (Time)"),
         T("Hora prevista de atencion.", "Expected time of attention.")],
    ], [1.1, 5.3]),
    ("p", T(
        "Esta tabla funciona como agenda de trabajo del dia: reune el dato de "
        "contacto y la ubicacion junto al horario, que es la informacion necesaria "
        "para organizar las visitas tecnicas.",
        "This table works as the day's work schedule: it brings together contact "
        "details and location alongside the time, which is the information needed "
        "to organize technical visits.")),

    # ----------------------------------------------------- 3.3 CLIENTES
    ("pagebreak",),
    ("h2", T("3.3 Clientes y solicitudes", "3.3 Clients and requests")),
    ("p", T(
        "Concentra la gestion completa de las solicitudes. A diferencia de la tabla "
        "de Inicio, aqui se listan todas las solicitudes sin importar su estado, e "
        "incorpora la barra de acciones que permite operar sobre ellas.",
        "This screen centralizes the full management of requests. Unlike the table "
        "on Inicio, here every request is listed regardless of its status, and it "
        "adds the action bar used to operate on them.")),
    ("figura", 3, T("Pantalla Clientes y solicitudes",
                    "The Clientes y solicitudes (Clients and requests) screen")),

    ("h3", T("3.3.1 Columnas de la tabla", "3.3.1 Table columns")),
    ("table", [
        [T("Columna", "Column"), T("Contenido", "Content")],
        [T("ID", "ID"), T("Numero de pedido de la solicitud.",
                          "Order number of the request.")],
        [T("Nombre", "Nombre (Name)"), T("Cliente solicitante.",
                                         "The requesting client.")],
        [T("Telefono", "Telefono (Phone)"), T("Numero de contacto.",
                                              "Contact number.")],
        [T("Direccion", "Direccion (Address)"), T("Domicilio del servicio.",
                                                  "Service address.")],
        [T("Estado", "Estado (Status)"),
         T("Etapa en la que se encuentra la solicitud: PENDIENTE, EN PROCESO o "
           "ATENDIDA.",
           "Stage the request is at: PENDIENTE (pending), EN PROCESO (in progress) "
           "or ATENDIDA (completed).")],
        [T("Fecha", "Fecha (Date)"), T("Dia asociado a la solicitud.",
                                       "Day associated with the request.")],
        [T("Hora", "Hora (Time)"), T("Hora asociada a la solicitud.",
                                     "Time associated with the request.")],
    ], [1.1, 5.3]),
    ("p", T(
        "El listado se agrupa por estado, presentando primero las solicitudes "
        "PENDIENTE, despues las que estan EN PROCESO y al final las ATENDIDA. Asi "
        "el trabajo por hacer queda siempre en la parte superior de la pantalla.",
        "The list is grouped by status, showing PENDIENTE requests first, then "
        "those EN PROCESO and finally the ATENDIDA ones. This keeps outstanding "
        "work always at the top of the screen.")),

    ("h3", T("3.3.2 Estados de una solicitud", "3.3.2 Request statuses")),
    ("p", T(
        "El estado describe el avance del servicio y es el eje del seguimiento:",
        "The status describes the progress of the service and is the backbone of "
        "tracking:")),
    ("table", [
        [T("Estado", "Status"), T("Significado", "Meaning")],
        [T("PENDIENTE", "PENDIENTE (pending)"),
         T("La solicitud fue recibida pero el taller aun no comienza a atenderla. "
           "Es el estado con el que nace toda solicitud enviada por un cliente.",
           "The request has been received but the workshop has not started working "
           "on it. This is the status every client-submitted request starts in.")],
        [T("EN PROCESO", "EN PROCESO (in progress)"),
         T("El servicio ya se esta ejecutando: se asigno un tecnico, se diagnostico "
           "el equipo o se esta realizando la reparacion.",
           "The service is already under way: a technician has been assigned, the "
           "equipment has been diagnosed or the repair is being carried out.")],
        [T("ATENDIDA", "ATENDIDA (completed)"),
         T("El servicio concluyo. La solicitud se conserva en el historial pero ya "
           "no requiere accion.",
           "The service has finished. The request is kept in the history but no "
           "longer requires action.")],
    ], [1.3, 5.1]),

    ("h3", T("3.3.3 Consultar el detalle de una solicitud",
             "3.3.3 Viewing the detail of a request")),
    ("p", T(
        "El texto guia bajo el titulo lo indica: al hacer doble clic sobre un "
        "cliente de la tabla se abre el detalle de su solicitud. Ahi puede "
        "consultarse la informacion completa del registro, incluido el motivo del "
        "servicio que el cliente describio al enviarlo, que por su extension no "
        "cabe en el listado.",
        "The helper text under the title says it: double-clicking a client in the "
        "table opens the detail of their request. There you can consult the full "
        "record, including the reason for the service the client described on "
        "submitting it, which is too long to fit in the list.")),

    ("h3", T("3.3.4 Barra de acciones", "3.3.4 Action bar")),
    ("p", T(
        "Los botones de la parte superior derecha operan sobre la solicitud "
        "previamente seleccionada en la tabla, con excepcion de Actualizar:",
        "The buttons at the top right act on the request previously selected in the "
        "table, with the exception of Actualizar:")),
    ("table", [
        [T("Boton", "Button"), T("Funcion", "Function")],
        [T("Cotizar", "Cotizar (Quote)"),
         T("Genera la cotizacion del servicio de la solicitud seleccionada, "
           "incorporando las piezas del inventario que se requieran y su precio.",
           "Generates the service quote for the selected request, incorporating the "
           "inventory parts required and their price.")],
        [T("Imprimir ticket", "Imprimir ticket (Print receipt)"),
         T("Emite el comprobante del servicio con los datos del cliente y los "
           "conceptos cobrados, para entregarlo como constancia.",
           "Issues the service receipt with the client's details and the items "
           "charged, to be handed over as proof.")],
        [T("Cambiar estado", "Cambiar estado (Change status)"),
         T("Hace avanzar la solicitud a la siguiente etapa, por ejemplo de "
           "PENDIENTE a EN PROCESO y despues a ATENDIDA.",
           "Advances the request to the next stage, for example from PENDIENTE to "
           "EN PROCESO and then to ATENDIDA.")],
        [T("Eliminar", "Eliminar (Delete)"),
         T("Borra del sistema la solicitud seleccionada. Se usa para descartar "
           "registros duplicados o capturados por error.",
           "Deletes the selected request from the system. Used to discard duplicate "
           "records or ones entered by mistake.")],
        [T("Actualizar", "Actualizar (Refresh)"),
         T("Vuelve a leer la informacion y refresca la tabla para mostrar las "
           "solicitudes nuevas que hayan llegado desde el modulo Cliente.",
           "Re-reads the information and refreshes the table to show new requests "
           "that have arrived from the Client Module.")],
    ], [1.3, 5.1]),
    ("nota", T(
        "Antes de pulsar Cotizar, Imprimir ticket, Cambiar estado o Eliminar es "
        "necesario seleccionar la fila del cliente correspondiente, ya que estas "
        "acciones se aplican al registro activo.",
        "Before clicking Cotizar, Imprimir ticket, Cambiar estado or Eliminar you "
        "must select the corresponding client row, since these actions apply to the "
        "active record.")),

    # --------------------------------------------------- 3.4 INVENTARIO
    ("pagebreak",),
    ("h2", T("3.4 Inventario de piezas", "3.4 Parts inventory")),
    ("p", T(
        "Es el catalogo de refacciones y equipos del taller. Cumple dos funciones: "
        "sirve como lista de precios al momento de cotizar un servicio y como "
        "control de existencias para saber de que material se dispone.",
        "This is the workshop's catalog of spare parts and equipment. It serves two "
        "purposes: as a price list when quoting a service, and as stock control to "
        "know what material is available.")),
    ("figura", 4, T("Pantalla Inventario de piezas",
                    "The Inventario de piezas (Parts inventory) screen")),

    ("h3", T("3.4.1 Columnas del catalogo", "3.4.1 Catalog columns")),
    ("table", [
        [T("Columna", "Column"), T("Contenido", "Content")],
        [T("Codigo", "Codigo (Code)"),
         T("Clave corta que identifica la pieza de forma unica. Se forma con un "
           "prefijo por familia y un sufijo por variante, por ejemplo BOM-100 para "
           "una bomba centrifuga de 1 HP o ARR-TER para un arrancador termico.",
           "Short key that uniquely identifies the part. It is formed from a family "
           "prefix and a variant suffix, for example BOM-100 for a 1 HP centrifugal "
           "pump or ARR-TER for a thermal starter.")],
        [T("Nombre", "Nombre (Name)"),
         T("Denominacion comercial de la pieza, normalmente con su capacidad o "
           "potencia, como Bomba centrifuga 1.5 HP.",
           "Commercial name of the part, usually including its capacity or power, "
           "such as Bomba centrifuga 1.5 HP.")],
        [T("Descripcion", "Descripcion (Description)"),
         T("Detalle tecnico complementario que precisa material, uso o "
           "caracteristicas del componente.",
           "Supplementary technical detail specifying the material, use or "
           "characteristics of the component.")],
        [T("Precio", "Precio (Price)"),
         T("Importe unitario de venta. Es el valor que se traslada a la cotizacion.",
           "Unit sale price. This is the value carried over to the quote.")],
        [T("Stock", "Stock"),
         T("Unidades disponibles en el almacen.", "Units available in the warehouse.")],
    ], [1.2, 5.2]),
    ("p", T(
        "El catalogo se presenta ordenado por nombre, lo que agrupa de forma natural "
        "las piezas de una misma familia y facilita localizarlas al recorrer la "
        "lista. Ejemplos de registros:",
        "The catalog is presented sorted by name, which naturally groups parts from "
        "the same family and makes them easier to find when scanning the list. "
        "Sample records:")),
    ("table", [
        [T("Codigo", "Codigo (Code)"), T("Nombre", "Nombre (Name)"),
         T("Precio", "Precio (Price)"), T("Stock", "Stock")],
        [T("ABR-MAN", "ABR-MAN"),
         T("Abrazadera para manguera", "Abrazadera para manguera (hose clamp)"),
         T("$18.00", "$18.00"), T("110", "110")],
        [T("ARR-TER", "ARR-TER"),
         T("Arrancador termico", "Arrancador termico (thermal starter)"),
         T("$540.00", "$540.00"), T("11", "11")],
        [T("BAT-100", "BAT-100"),
         T("Bateria ciclo profundo", "Bateria ciclo profundo (deep cycle battery)"),
         T("$4,300.00", "$4,300.00"), T("5", "5")],
        [T("BOM-100", "BOM-100"),
         T("Bomba centrifuga 1 HP", "Bomba centrifuga 1 HP (1 HP centrifugal pump)"),
         T("$2,450.00", "$2,450.00"), T("4", "4")],
        [T("BOM-SOL2", "BOM-SOL2"),
         T("Bomba solar sumergible",
           "Bomba solar sumergible (solar submersible pump)"),
         T("$14,500.00", "$14,500.00"), T("1", "1")],
    ], [1.1, 3.1, 1.2, 1.0]),

    ("h3", T("3.4.2 Acciones sobre el inventario", "3.4.2 Inventory actions")),
    ("table", [
        [T("Boton", "Button"), T("Funcion", "Function")],
        [T("+ Agregar", "+ Agregar (Add)"),
         T("Da de alta una pieza nueva en el catalogo. Requiere definir su codigo, "
           "nombre, descripcion, precio y existencia inicial.",
           "Registers a new part in the catalog. It requires defining its code, "
           "name, description, price and initial stock.")],
        [T("Editar", "Editar (Edit)"),
         T("Modifica los datos de la pieza seleccionada. Es la via para actualizar "
           "un precio o corregir la cantidad en stock.",
           "Modifies the data of the selected part. This is the way to update a "
           "price or correct the quantity in stock.")],
        [T("Eliminar", "Eliminar (Delete)"),
         T("Retira del catalogo la pieza seleccionada.",
           "Removes the selected part from the catalog.")],
    ], [1.3, 5.1]),
    ("nota", T(
        "Mantener el stock y los precios al dia es lo que garantiza que las "
        "cotizaciones y los reportes reflejen la realidad del taller, ya que ambos "
        "se alimentan de este catalogo.",
        "Keeping stock and prices up to date is what guarantees that quotes and "
        "reports reflect the reality of the workshop, since both are fed by this "
        "catalog.")),

    # ----------------------------------------------------- 3.5 REPORTES
    ("pagebreak",),
    ("h2", T("3.5 Reportes y estadisticas", "3.5 Reports and statistics")),
    ("p", T(
        "Reune la informacion de analisis del taller. Combina cuatro indicadores de "
        "ventas con tres tablas que responden a preguntas concretas: que se vende "
        "mas, como evolucionan las ventas y que material esta por agotarse.",
        "This screen gathers the workshop's analytical information. It combines "
        "four sales indicators with three tables that answer specific questions: "
        "what sells most, how sales are trending and which material is about to run "
        "out.")),
    ("figura", 5, T("Pantalla Reportes y estadisticas",
                    "The Reportes y estadisticas (Reports and statistics) screen")),

    ("h3", T("3.5.1 Indicadores de venta", "3.5.1 Sales indicators")),
    ("table", [
        [T("Indicador", "Indicator"), T("Significado", "Meaning"),
         T("Ejemplo", "Example")],
        [T("Ventas totales", "Ventas totales (Total sales)"),
         T("Numero de ventas registradas desde el inicio de operacion.",
           "Number of sales recorded since operations began."), T("12", "12")],
        [T("Ingresos totales", "Ingresos totales (Total revenue)"),
         T("Importe acumulado de todas esas ventas.",
           "Accumulated amount of all those sales."),
         T("$16,735.00", "$16,735.00")],
        [T("Ventas de la semana", "Ventas de la semana (Sales this week)"),
         T("Importe vendido en la semana en curso y, entre parentesis, el numero de "
           "operaciones.",
           "Amount sold in the current week and, in parentheses, the number of "
           "transactions."), T("$0.00 (0)", "$0.00 (0)")],
        [T("Ventas del mes", "Ventas del mes (Sales this month)"),
         T("Importe vendido en el mes en curso y el numero de operaciones "
           "correspondiente.",
           "Amount sold in the current month and the corresponding number of "
           "transactions."), T("$10,785.00 (7)", "$10,785.00 (7)")],
    ], [1.5, 3.6, 1.3]),
    ("p", T(
        "Los indicadores semanal y mensual permiten comparar el ritmo reciente con "
        "el acumulado. Un valor en cero en la semana, con ventas presentes en el "
        "mes, senala simplemente que en los ultimos dias no se ha registrado "
        "ninguna operacion.",
        "The weekly and monthly indicators allow the recent pace to be compared "
        "with the cumulative total. A value of zero for the week, with sales "
        "present for the month, simply signals that no transaction has been "
        "recorded in the last few days.")),

    ("h3", T("3.5.2 Piezas mas vendidas", "3.5.2 Best-selling parts")),
    ("p", T(
        "Ordena las piezas segun su rotacion. Presenta el nombre de la Pieza, las "
        "Unidades vendidas y los Ingresos que ha generado cada una. Es la "
        "referencia para decidir que material conviene mantener siempre disponible.",
        "Ranks parts by turnover. It shows the Pieza (part) name, the Unidades "
        "vendidas (units sold) and the Ingresos (revenue) each one has generated. "
        "It is the reference for deciding which material is worth always keeping "
        "available.")),
    ("table", [
        [T("Pieza", "Pieza (Part)"),
         T("Unidades vendidas", "Unidades vendidas (Units sold)"),
         T("Ingresos", "Ingresos (Revenue)")],
        [T("Impulsor 1/2 HP", "Impulsor 1/2 HP (1/2 HP impeller)"),
         T("2", "2"), T("$460.00", "$460.00")],
        [T("Presostato automatico",
           "Presostato automatico (automatic pressure switch)"),
         T("2", "2"), T("$360.00", "$360.00")],
        [T("Cinta teflon", "Cinta teflon (Teflon tape)"),
         T("2", "2"), T("$30.00", "$30.00")],
    ], [3.0, 1.9, 1.5]),
    ("p", T(
        "Comparar unidades contra ingresos distingue las piezas de alta rotacion y "
        "bajo importe de aquellas que, con pocas unidades, aportan mas dinero.",
        "Comparing units against revenue distinguishes high-turnover, low-value "
        "parts from those that, with few units, bring in more money.")),

    ("h3", T("3.5.3 Ventas por mes", "3.5.3 Sales by month")),
    ("p", T(
        "Agrupa las ventas por periodo mensual, mostrando el Mes en formato "
        "ano-mes, el N. de ventas y el Total facturado. Permite observar la "
        "tendencia del negocio a lo largo del tiempo.",
        "Groups sales by monthly period, showing the Mes (month) in year-month "
        "format, the N. de ventas (number of sales) and the Total invoiced. It "
        "makes it possible to observe the trend of the business over time.")),
    ("table", [
        [T("Mes", "Mes (Month)"), T("N. de ventas", "N. de ventas (No. of sales)"),
         T("Total", "Total")],
        [T("2026-07", "2026-07"), T("7", "7"), T("$10,785.00", "$10,785.00")],
        [T("2026-06", "2026-06"), T("2", "2"), T("$3,455.00", "$3,455.00")],
    ], [2.0, 2.2, 2.2]),

    ("h3", T("3.5.4 Piezas con stock bajo", "3.5.4 Parts with low stock")),
    ("p", T(
        "Lista unicamente las piezas cuya existencia es igual o menor a 10 "
        "unidades, umbral indicado en el propio titulo de la seccion. Muestra el "
        "Codigo, la Pieza y el Stock restante. Funciona como alerta de "
        "reabastecimiento: las piezas que aparecen aqui son las que deben pedirse "
        "al proveedor antes de quedarse sin material.",
        "Lists only the parts whose stock is equal to or below 10 units, the "
        "threshold stated in the section title itself. It shows the Codigo (code), "
        "the Pieza (part) and the remaining Stock. It works as a restocking alert: "
        "the parts appearing here are the ones to order from the supplier before "
        "running out of material.")),
    ("table", [
        [T("Codigo", "Codigo (Code)"), T("Pieza", "Pieza (Part)"),
         T("Stock", "Stock")],
        [T("MOT-750", "MOT-750"),
         T("Motor 7.5 HP", "Motor 7.5 HP (7.5 HP motor)"), T("1", "1")],
        [T("BOM-SUM5", "BOM-SUM5"),
         T("Bomba sumergible 5 HP",
           "Bomba sumergible 5 HP (5 HP submersible pump)"), T("1", "1")],
    ], [1.5, 3.4, 1.5]),

    ("h3", T("3.5.5 Actualizar", "3.5.5 Refresh")),
    ("p", T(
        "El boton Actualizar de la esquina superior derecha recalcula todos los "
        "indicadores y tablas de la pantalla. Conviene pulsarlo despues de "
        "registrar ventas o de modificar el inventario para consultar cifras "
        "vigentes.",
        "The Actualizar (Refresh) button in the top right corner recalculates every "
        "indicator and table on the screen. It is worth clicking after recording "
        "sales or modifying the inventory in order to see current figures.")),

    # -------------------------------------------------------- 4. FLUJO
    ("pagebreak",),
    ("h1", T("4. Flujo de trabajo completo", "4. Complete workflow")),
    ("p", T(
        "Las pantallas descritas se articulan en un recorrido unico, desde que el "
        "cliente reporta la falla hasta que el servicio queda cerrado y reflejado "
        "en los reportes:",
        "The screens described come together in a single journey, from the moment "
        "the client reports the fault until the service is closed and reflected in "
        "the reports:")),
    ("table", [
        [T("Paso", "Step"), T("Pantalla", "Screen"), T("Que ocurre", "What happens")],
        [T("1", "1"), T("Cliente / Solicitar servicio",
                        "Client / Solicitar servicio"),
         T("El cliente captura nombre, telefono, direccion y motivo, y envia la "
           "solicitud. Recibe su numero de pedido.",
           "The client enters name, phone, address and reason, and submits the "
           "request. They receive their order number.")],
        [T("2", "2"), T("Administrador / Inicio", "Administrator / Inicio"),
         T("La solicitud aparece como PENDIENTE en Proximos servicios y suma al "
           "contador de pendientes.",
           "The request appears as PENDIENTE under Proximos servicios and adds to "
           "the pending counter.")],
        [T("3", "3"), T("Administrador / Clientes", "Administrator / Clientes"),
         T("El administrador abre el detalle con doble clic, revisa el motivo del "
           "servicio y genera la cotizacion con Cotizar.",
           "The administrator opens the detail with a double click, reviews the "
           "reason for the service and generates the quote with Cotizar.")],
        [T("4", "4"), T("Administrador / Inventario", "Administrator / Inventario"),
         T("Se consultan precios y existencias de las piezas necesarias y se ajusta "
           "el stock con Editar.",
           "Prices and stock of the required parts are checked and stock is "
           "adjusted with Editar.")],
        [T("5", "5"), T("Administrador / Clientes", "Administrator / Clientes"),
         T("Con Cambiar estado la solicitud pasa a EN PROCESO mientras se realiza "
           "el trabajo, y a ATENDIDA al concluirlo. Con Imprimir ticket se entrega "
           "el comprobante.",
           "With Cambiar estado the request moves to EN PROCESO while the work is "
           "carried out, and to ATENDIDA on completion. With Imprimir ticket the "
           "receipt is handed over.")],
        [T("6", "6"), T("Administrador / Reportes", "Administrator / Reportes"),
         T("La operacion se incorpora a los indicadores de ventas, a las piezas mas "
           "vendidas y a las alertas de stock bajo.",
           "The transaction is incorporated into the sales indicators, the "
           "best-selling parts and the low stock alerts.")],
    ], [0.55, 1.85, 4.0]),

    # ----------------------------------------------------- 5. GLOSARIO
    ("h1", T("5. Glosario", "5. Glossary")),
    ("table", [
        [T("Termino", "Term"), T("Definicion", "Definition")],
        [T("Solicitud de servicio", "Solicitud de servicio (service request)"),
         T("Peticion registrada por un cliente para que el taller revise o repare "
           "una bomba de agua.",
           "A request logged by a client for the workshop to inspect or repair a "
           "water pump.")],
        [T("Numero de pedido", "Numero de pedido (order number)"),
         T("Identificador consecutivo que el sistema asigna a cada solicitud al "
           "enviarse. Corresponde a la columna ID.",
           "Sequential identifier the system assigns to each request on "
           "submission. It corresponds to the ID column.")],
        [T("Estado", "Estado (status)"),
         T("Etapa de avance de una solicitud: PENDIENTE, EN PROCESO o ATENDIDA.",
           "Progress stage of a request: PENDIENTE, EN PROCESO or ATENDIDA.")],
        [T("Codigo de pieza", "Codigo de pieza (part code)"),
         T("Clave corta y unica que identifica una refaccion en el inventario, como "
           "BOM-100.",
           "Short, unique key identifying a spare part in the inventory, such as "
           "BOM-100.")],
        [T("Stock", "Stock"),
         T("Numero de unidades de una pieza disponibles en el almacen.",
           "Number of units of a part available in the warehouse.")],
        [T("Stock bajo", "Stock bajo (low stock)"),
         T("Condicion de una pieza cuya existencia es menor o igual a 10 unidades; "
           "el sistema la reporta como alerta de reabastecimiento.",
           "Condition of a part whose stock is less than or equal to 10 units; the "
           "system reports it as a restocking alert.")],
        [T("Cotizacion", "Cotizacion (quote)"),
         T("Documento que estima el costo de un servicio a partir de las piezas y "
           "su precio de catalogo.",
           "Document estimating the cost of a service based on the parts and their "
           "catalog price.")],
        [T("Ticket", "Ticket (receipt)"),
         T("Comprobante impreso del servicio realizado y de los conceptos cobrados.",
           "Printed proof of the service performed and the items charged.")],
    ], [1.6, 4.8]),
]


# ============================================================== renderer
def render(doc, lang):
    """Escribe todos los bloques en el idioma indicado."""
    for bloque in BLOQUES:
        tipo = bloque[0]

        if tipo == "pagebreak":
            doc.add_page_break()

        elif tipo in ("h1", "h2", "h3"):
            doc.add_heading(bloque[1][lang], level=int(tipo[1]))

        elif tipo == "p":
            para(doc, bloque[1][lang])

        elif tipo == "bullets":
            for item in bloque[1]:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(3)
                if "b" in item:
                    r = p.add_run(item["b"][lang])
                    r.bold = True
                    r.font.size = Pt(10.5)
                    r2 = p.add_run(" " + item["t"][lang])
                    r2.font.size = Pt(10.5)
                else:
                    p.add_run(item["t"][lang]).font.size = Pt(10.5)

        elif tipo == "steps":
            for item in bloque[1]:
                p = doc.add_paragraph(style="List Number")
                p.paragraph_format.space_after = Pt(3)
                p.add_run(item[lang]).font.size = Pt(10.5)

        elif tipo == "table":
            filas = [[celda[lang] for celda in fila] for fila in bloque[1]]
            anchos = bloque[2] if len(bloque) > 2 else None
            data_table(doc, filas, anchos)
            para(doc, "", space_after=2)

        elif tipo == "figura":
            figura(doc, bloque[1], bloque[2][lang], lang)

        elif tipo == "nota":
            nota(doc, bloque[1][lang], lang)

        else:
            raise ValueError(f"Bloque desconocido: {tipo}")


def portada(doc, lang, es_cubierta):
    """Portada del documento (lang=ES) o separador de la version en ingles."""
    for _ in range(4 if es_cubierta else 6):
        doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_after = Pt(0)
    r = t.add_run("VORTEX")
    r.font.size = Pt(58)
    r.bold = True
    r.font.color.rgb = BLUE

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.paragraph_format.space_after = Pt(24)
    r = st.add_run("Sistema de Gestion para Taller de Bombas de Agua" if lang == ES
                   else "Water Pump Workshop Management System")
    r.font.size = Pt(13)
    r.font.color.rgb = GRAY

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    horizontal_rule(rule, BLUE_HEX, 24)
    rule.paragraph_format.space_after = Pt(22)

    mt = doc.add_paragraph()
    mt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mt.paragraph_format.space_after = Pt(6)
    r = mt.add_run("MANUAL DE USUARIO" if lang == ES else "USER MANUAL")
    r.font.size = Pt(26)
    r.bold = True
    r.font.color.rgb = NAVY

    sb = doc.add_paragraph()
    sb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sb.paragraph_format.space_after = Pt(10)
    r = sb.add_run("Guia de funcionamiento de los modulos Cliente y Administrador"
                   if lang == ES else
                   "A guide to how the Client and Administrator modules work")
    r.font.size = Pt(11)
    r.italic = True
    r.font.color.rgb = GRAY

    idioma = doc.add_paragraph()
    idioma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = idioma.add_run("Version en espanol  ·  English version included"
                       if lang == ES else
                       "English version  ·  Version en espanol incluida")
    r.font.size = Pt(9.5)
    r.font.color.rgb = BLUE

    for _ in range(6 if es_cubierta else 3):
        doc.add_paragraph()

    etiquetas = ([("Documento", "Manual de usuario"), ("Version", "1.0"),
                  ("Fecha", "Julio de 2026")] if lang == ES else
                 [("Document", "User manual"), ("Version", "1.0"),
                  ("Date", "July 2026")])
    info = doc.add_table(rows=3, cols=2)
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, (k, v) in enumerate(etiquetas):
        c0, c1 = info.cell(i, 0), info.cell(i, 1)
        c0.width, c1.width = Inches(1.5), Inches(2.6)
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rk = c0.paragraphs[0].add_run(k)
        rk.bold = True
        rk.font.size = Pt(10)
        rk.font.color.rgb = NAVY
        rv = c1.paragraphs[0].add_run("   " + v)
        rv.font.size = Pt(10)
        rv.font.color.rgb = GRAY


# ================================================================ armado
doc = Document()
doc.core_properties.title = "Manual de Usuario / User Manual - VORTEX"
doc.core_properties.subject = ("Sistema de Gestion para Taller de Bombas de Agua / "
                               "Water Pump Workshop Management System")
doc.core_properties.author = "VORTEX"
doc.core_properties.language = "es-MX, en-US"

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

for nombre, tam, color in (("Heading 1", 17, NAVY), ("Heading 2", 13.5, BLUE),
                           ("Heading 3", 11.5, NAVY)):
    s = doc.styles[nombre]
    s.font.name = "Calibri"
    s.font.size = Pt(tam)
    s.font.bold = True
    s.font.color.rgb = color
    s.paragraph_format.space_before = Pt(14 if nombre == "Heading 1" else 10)
    s.paragraph_format.space_after = Pt(6)
    s.paragraph_format.keep_with_next = True

for s in doc.sections:
    s.top_margin = Inches(1.0)
    s.bottom_margin = Inches(0.9)
    s.left_margin = Inches(1.0)
    s.right_margin = Inches(1.0)

# ---- seccion 1: portada, sin numeracion
portada(doc, ES, es_cubierta=True)

# ---- seccion 2: todo el contenido, numeracion corrida desde 1
cuerpo = doc.add_section(WD_SECTION.NEW_PAGE)
cuerpo.top_margin = Inches(1.0)
cuerpo.bottom_margin = Inches(0.9)
cuerpo.left_margin = Inches(1.0)
cuerpo.right_margin = Inches(1.0)
restart_page_numbering(cuerpo, 1)

cuerpo.header.is_linked_to_previous = False
hp = cuerpo.header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = hp.add_run("VORTEX")
r.bold = True
r.font.size = Pt(10)
r.font.color.rgb = BLUE
r2 = hp.add_run("  |  Manual de Usuario  ·  User Manual")
r2.font.size = Pt(9)
r2.font.color.rgb = GRAY
horizontal_rule(hp, BORDER_HEX, 6)

cuerpo.footer.is_linked_to_previous = False
fp = cuerpo.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER


def texto_pie(txt):
    r = fp.add_run(txt)
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY


texto_pie("Pagina ")
add_field(fp, "PAGE")
texto_pie(" de ")
add_field(fp, "SECTIONPAGES")
texto_pie("        ·        Page ")
add_field(fp, "PAGE")
texto_pie(" of ")
add_field(fp, "SECTIONPAGES")

# ---------------------------------------------- version en espanol
h = doc.add_heading("Contenido", level=1)
add_toc(doc, "VERSION_ES")
doc.add_page_break()

ancla_es = doc.add_paragraph()
ancla_es.paragraph_format.space_after = Pt(0)
bid_es = bookmark_start(ancla_es, "VERSION_ES")
render(doc, ES)
bookmark_end(doc.paragraphs[-1], bid_es)

# ---------------------------------------------- version en ingles
doc.add_page_break()
portada(doc, EN, es_cubierta=False)
doc.add_page_break()

doc.add_heading("Contents", level=1)
add_toc(doc, "VERSION_EN")
doc.add_page_break()

ancla_en = doc.add_paragraph()
ancla_en.paragraph_format.space_after = Pt(0)
bid_en = bookmark_start(ancla_en, "VERSION_EN")
render(doc, EN)
bookmark_end(doc.paragraphs[-1], bid_en)

# ------------------------------------------------------------- guardar
salida = os.path.join(BASE, "Manual_de_Usuario_VORTEX.docx")
doc.save(salida)

print("Documento generado:", salida)
print()
titulos = {
    1: "Cliente - Solicitar servicio",
    2: "Administrador - Inicio",
    3: "Administrador - Clientes y solicitudes",
    4: "Administrador - Inventario de piezas",
    5: "Administrador - Reportes y estadisticas",
}
print("Asignacion de capturas:")
for num in sorted(MAPEO):
    print(f"  Figura/Figure {num}  {titulos[num]:<42} <-  {MAPEO[num]}")
faltan = [n for n, a in MAPEO.items() if a.startswith("--")]
if faltan:
    print()
    print(f"Figuras sin imagen: {faltan}")
    print("Coloque figura1.png ... figura5.png en ./capturas/ y ejecute de nuevo.")
